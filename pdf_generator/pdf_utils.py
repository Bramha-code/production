import re
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def generate_docx_report(text_input: str, template_path: str, component_name: str):
    print(f"[DOCX] Starting generation using template: {template_path}")
    
    doc = Document(template_path) if os.path.exists(template_path) else Document()

    # 1. Global Replacement (Header/Footer/Body)
    target_tag = "{{COMPONENT_NAME}}"
    display_name = component_name.replace('_', ' ')
    _global_replace(target_tag, display_name, doc)

    # 2. Table of Contents
    doc.add_page_break()
    doc.add_heading('Table of Contents', level=1)
    # Add a visible placeholder so the user knows to right-click
    p = doc.add_paragraph("Right-click here and select 'Update Field' to generate page numbers.")
    _add_word_toc(doc)
    doc.add_page_break()

    # 3. Clean and Parse Content
    # Remove LLM meta-talk (End of Chapter, Next: Chapter, etc.)
    clean_text = re.sub(r'(?i)---.*?(End of Chapter|Next:).*?---', '', text_input, flags=re.DOTALL)
    clean_text = re.sub(r'(?i)> Note:.*?\n', '', clean_text) # Remove the blockquote notes
    
    pattern = r'(^\s*\|.*\|(?:\n\s*\|.*\|)*)|(^\s*#{1,6}\s+.*$)'
    tokens = re.split(pattern, clean_text, flags=re.MULTILINE)

    for token in tokens:
        if not token or not token.strip(): continue
        token = token.strip()

        if token.startswith('|'):
            _add_markdown_table_to_doc(doc, token)
        elif token.startswith('#'):
            level = len(re.match(r'#+', token).group())
            clean_title = token.lstrip('#').strip()
            
            # FEATURE: Always start Level 1 headers on a new page
            if level == 1:
                doc.add_page_break()
            
            if clean_title:
                doc.add_heading(clean_title, level=min(level, 9))
        else:
            # Handle manual page break markers if any remain
            if re.match(r'^-{3,}$', token):
                doc.add_page_break()
            else:
                _add_formatted_paragraph(doc, token)

    # 4. Save
    output_path = f"{component_name.replace(' ', '_')}_report.docx"
    doc.save(output_path)
    print(f"[DOCX] Success! File saved: {output_path}")
    return output_path

def _add_word_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2); run._r.append(fldChar3)

def _global_replace(target, replacement, doc_obj):
    def replace_in_p(p):
        if target in p.text:
            new_text = p.text.replace(target, replacement)
            for run in p.runs: run.text = ""
            if p.runs: p.runs[0].text = new_text
            else: p.add_run(new_text)
    for p in doc_obj.paragraphs: replace_in_p(p)
    for s in doc_obj.sections:
        for hf in [s.header, s.footer]:
            for p in hf.paragraphs: replace_in_p(p)
    for t in doc_obj.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs: replace_in_p(p)

def _add_markdown_table_to_doc(doc, table_str):
    rows = [r.strip() for r in table_str.split('\n') if r.strip()]
    data_rows = [r for r in rows if not re.match(r'^\|?[\s\-:|]+\|?$', r)]
    if not data_rows: return
    table = doc.add_table(rows=len(data_rows), cols=max(row.count('|') for row in data_rows) - 1)
    table.style = 'Table Grid'
    for i, row_str in enumerate(data_rows):
        cells = [c.strip() for c in row_str.strip('|').split('|')]
        for j, val in enumerate(cells):
            if j < len(table.rows[i].cells):
                _process_inline_bold(table.rows[i].cells[j].paragraphs[0], val)

def _add_formatted_paragraph(doc, text):
    for line in text.split('\n'):
        if line.strip(): _process_inline_bold(doc.add_paragraph(), line)

def _process_inline_bold(p, text):
    parts = re.split(r'(\*\*.*?\*\*|<br\s*/?>|<sub>.*?</sub>)', text, flags=re.IGNORECASE)
    for part in parts:
        if not part: continue
        if part.startswith('**'): run = p.add_run(part[2:-2]); run.bold = True
        elif part.lower().startswith('<sub>'): run = p.add_run(part[5:-6]); run.font.subscript = True
        elif re.match(r'<br\s*/?>', part, re.IGNORECASE): p.add_run().add_break()
        else: p.add_run(part)