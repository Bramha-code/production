"""
Document Exporter - DOCX-based Document Generation

This module provides document export functionality using python-docx.
It replaces the ReportLab PDF generation with Microsoft Word document generation.

Features:
- Markdown to DOCX conversion
- Table of Contents generation
- Professional formatting with templates
- Test plan export with structured sections
"""

import re
import os
import io
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX export will not work.")


# =========================================================
# Document Exporter Class
# =========================================================

class DocumentExporter:
    """
    Export content to DOCX format using python-docx.

    Supports:
    - Markdown-like content conversion
    - Template-based document generation
    - Test plan structured export
    - Tables and formatting
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize the document exporter.

        Args:
            template_path: Optional path to a DOCX template file
        """
        self.template_path = template_path

        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    def export_to_docx(self, content: str, title: str = "Document",
                       component_name: Optional[str] = None) -> bytes:
        """
        Export markdown-like content to DOCX format.

        Args:
            content: The content to export (markdown-like format)
            title: Document title
            component_name: Optional component name for template replacement

        Returns:
            DOCX document as bytes
        """
        # Create document from template or blank
        if self.template_path and os.path.exists(self.template_path):
            doc = Document(self.template_path)
        else:
            doc = Document()

        # Global replacement if component name provided
        if component_name:
            target_tag = "{{COMPONENT_NAME}}"
            display_name = component_name.replace('_', ' ')
            self._global_replace(target_tag, display_name, doc)

        # Add title
        doc.add_heading(title, level=0)

        # Add Table of Contents
        doc.add_page_break()
        doc.add_heading('Table of Contents', level=1)
        p = doc.add_paragraph("Right-click and select 'Update Field' to generate page numbers.")
        self._add_word_toc(doc)
        doc.add_page_break()

        # Clean and parse content
        clean_content = self._clean_content(content)
        self._parse_and_add_content(doc, clean_content)

        # Return as bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_test_plan_to_docx(self, test_plan: Any) -> bytes:
        """
        Export a TestPlan object to DOCX format.

        Args:
            test_plan: TestPlan object with structured data

        Returns:
            DOCX document as bytes
        """
        doc = Document()

        # Title page
        title = test_plan.title if hasattr(test_plan, 'title') else "EMC Test Plan"
        doc.add_heading(title, level=0)

        # Metadata
        if hasattr(test_plan, 'document_number'):
            doc.add_paragraph(f"Document Number: {test_plan.document_number}")
        if hasattr(test_plan, 'revision'):
            doc.add_paragraph(f"Revision: {test_plan.revision}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        doc.add_page_break()

        # Table of Contents
        doc.add_heading('Table of Contents', level=1)
        self._add_word_toc(doc)
        doc.add_page_break()

        # 1. Scope
        doc.add_heading('1. Scope', level=1)
        if hasattr(test_plan, 'scope') and test_plan.scope:
            doc.add_paragraph(test_plan.scope)
        else:
            doc.add_paragraph("This test plan defines the EMC testing requirements and procedures.")

        # 2. Applicable Standards
        doc.add_heading('2. Applicable Standards', level=1)
        if hasattr(test_plan, 'applicable_standards') and test_plan.applicable_standards:
            for std in test_plan.applicable_standards:
                doc.add_paragraph(f"• {std}", style='List Bullet')
        else:
            doc.add_paragraph("No applicable standards specified.")

        # 3. Equipment Under Test
        doc.add_heading('3. Equipment Under Test (EUT)', level=1)
        if hasattr(test_plan, 'eut_description') and test_plan.eut_description:
            doc.add_paragraph(test_plan.eut_description)
        else:
            doc.add_paragraph("EUT description not specified.")

        # 4. Test Equipment Required
        doc.add_heading('4. Test Equipment Required', level=1)
        if hasattr(test_plan, 'all_equipment') and test_plan.all_equipment:
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Equipment'
            hdr_cells[1].text = 'Source'

            for eq in test_plan.all_equipment:
                row_cells = table.add_row().cells
                if isinstance(eq, dict):
                    row_cells[0].text = eq.get('name', str(eq))
                    row_cells[1].text = eq.get('source', 'N/A')
                else:
                    row_cells[0].text = str(eq)
                    row_cells[1].text = 'N/A'
        else:
            doc.add_paragraph("No equipment specified.")

        # 5. Environmental Conditions
        doc.add_heading('5. Environmental Conditions', level=1)
        if hasattr(test_plan, 'environmental_conditions') and test_plan.environmental_conditions:
            for key, value in test_plan.environmental_conditions.items():
                doc.add_paragraph(f"• {key}: {value}", style='List Bullet')
        else:
            doc.add_paragraph("Standard laboratory conditions apply.")

        # 6. Test Cases
        doc.add_heading('6. Test Cases', level=1)
        if hasattr(test_plan, 'test_cases') and test_plan.test_cases:
            for i, tc in enumerate(test_plan.test_cases, 1):
                tc_id = tc.test_case_id if hasattr(tc, 'test_case_id') else f"TC-{i:03d}"
                tc_name = tc.objective if hasattr(tc, 'objective') else f"Test Case {i}"

                doc.add_heading(f"6.{i} {tc_id}: {tc_name}", level=2)

                # Test type and priority
                if hasattr(tc, 'test_type'):
                    doc.add_paragraph(f"Test Type: {tc.test_type}")
                if hasattr(tc, 'priority'):
                    doc.add_paragraph(f"Priority: {tc.priority}")

                # Requirement
                if hasattr(tc, 'requirement_text') and tc.requirement_text:
                    doc.add_heading("Requirement:", level=3)
                    doc.add_paragraph(tc.requirement_text)

                # Pre-conditions
                if hasattr(tc, 'pre_conditions') and tc.pre_conditions:
                    doc.add_heading("Pre-conditions:", level=3)
                    for pre in tc.pre_conditions:
                        doc.add_paragraph(f"• {pre}", style='List Bullet')

                # Procedure
                if hasattr(tc, 'procedure_steps') and tc.procedure_steps:
                    doc.add_heading("Procedure:", level=3)
                    for j, step in enumerate(tc.procedure_steps, 1):
                        doc.add_paragraph(f"{j}. {step}")

                # Pass/Fail Criteria
                if hasattr(tc, 'pass_fail_criteria') and tc.pass_fail_criteria:
                    doc.add_heading("Pass/Fail Criteria:", level=3)
                    doc.add_paragraph(tc.pass_fail_criteria)

                # Test Limits
                if hasattr(tc, 'test_limits') and tc.test_limits:
                    doc.add_heading("Test Limits:", level=3)
                    for limit_key, limit_val in tc.test_limits.items():
                        doc.add_paragraph(f"• {limit_key}: {limit_val}", style='List Bullet')
        else:
            doc.add_paragraph("No test cases defined.")

        # 7. Requirement Coverage Matrix
        doc.add_heading('7. Requirement Coverage Matrix', level=1)
        if hasattr(test_plan, 'coverage_matrix') and test_plan.coverage_matrix:
            coverage = test_plan.coverage_matrix
            if hasattr(coverage, 'coverage_percentage'):
                doc.add_paragraph(f"Overall Coverage: {coverage.coverage_percentage:.1f}%")

            if hasattr(coverage, 'items') and coverage.items:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Requirement ID'
                hdr_cells[1].text = 'Test Cases'
                hdr_cells[2].text = 'Status'

                for item in coverage.items[:20]:  # Limit to 20 rows
                    row_cells = table.add_row().cells
                    row_cells[0].text = item.requirement_id if hasattr(item, 'requirement_id') else str(item)
                    row_cells[1].text = ', '.join(item.test_case_ids) if hasattr(item, 'test_case_ids') else 'N/A'
                    row_cells[2].text = item.status if hasattr(item, 'status') else 'N/A'
        else:
            doc.add_paragraph("Coverage matrix not available.")

        # 8. Approval Signatures
        doc.add_heading('8. Approval', level=1)
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        headers = ['Role', 'Name', 'Signature / Date']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header

        roles = ['Prepared By:', 'Reviewed By:', 'Approved By:']
        for i, role in enumerate(roles, 1):
            table.rows[i].cells[0].text = role

        # Return as bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _clean_content(self, content: str) -> str:
        """Clean LLM meta-talk and formatting from content."""
        # Remove LLM meta-talk
        clean = re.sub(r'(?i)---.*?(End of Chapter|Next:).*?---', '', content, flags=re.DOTALL)
        clean = re.sub(r'(?i)> Note:.*?\n', '', clean)
        return clean

    def _parse_and_add_content(self, doc: Document, content: str):
        """Parse markdown-like content and add to document."""
        # Split into tokens (tables and headers vs other content)
        pattern = r'(^\s*\|.*\|(?:\n\s*\|.*\|)*)|(^\s*#{1,6}\s+.*$)'
        tokens = re.split(pattern, content, flags=re.MULTILINE)

        for token in tokens:
            if not token or not token.strip():
                continue

            token = token.strip()

            if token.startswith('|'):
                # Table
                self._add_markdown_table(doc, token)
            elif token.startswith('#'):
                # Header
                level = len(re.match(r'#+', token).group())
                clean_title = token.lstrip('#').strip()

                # Page break before level 1 headers
                if level == 1:
                    doc.add_page_break()

                if clean_title:
                    doc.add_heading(clean_title, level=min(level, 9))
            elif re.match(r'^-{3,}$', token):
                # Page break marker
                doc.add_page_break()
            else:
                # Regular paragraph
                self._add_formatted_paragraph(doc, token)

    def _add_word_toc(self, doc: Document):
        """Add a Word Table of Contents field."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def _global_replace(self, target: str, replacement: str, doc: Document):
        """Replace target text throughout the document."""
        def replace_in_paragraph(p):
            if target in p.text:
                new_text = p.text.replace(target, replacement)
                for run in p.runs:
                    run.text = ""
                if p.runs:
                    p.runs[0].text = new_text
                else:
                    p.add_run(new_text)

        # Replace in paragraphs
        for p in doc.paragraphs:
            replace_in_paragraph(p)

        # Replace in headers/footers
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                for p in hf.paragraphs:
                    replace_in_paragraph(p)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)

    def _add_markdown_table(self, doc: Document, table_str: str):
        """Convert markdown table to Word table."""
        rows = [r.strip() for r in table_str.split('\n') if r.strip()]

        # Filter out separator rows (|---|---|)
        data_rows = [r for r in rows if not re.match(r'^\|?[\s\-:|]+\|?$', r)]

        if not data_rows:
            return

        # Determine number of columns
        num_cols = max(row.count('|') for row in data_rows) - 1
        if num_cols <= 0:
            return

        table = doc.add_table(rows=len(data_rows), cols=num_cols)
        table.style = 'Table Grid'

        for i, row_str in enumerate(data_rows):
            cells = [c.strip() for c in row_str.strip('|').split('|')]
            for j, val in enumerate(cells):
                if j < len(table.rows[i].cells):
                    self._add_formatted_text(table.rows[i].cells[j].paragraphs[0], val)

    def _add_formatted_paragraph(self, doc: Document, text: str):
        """Add paragraph with inline formatting."""
        for line in text.split('\n'):
            if line.strip():
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with inline formatting (bold, subscript, line breaks)."""
        # Split on formatting markers
        parts = re.split(r'(\*\*.*?\*\*|<br\s*/?>|<sub>.*?</sub>)', text, flags=re.IGNORECASE)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.lower().startswith('<sub>') and part.lower().endswith('</sub>'):
                # Subscript
                run = paragraph.add_run(part[5:-6])
                run.font.subscript = True
            elif re.match(r'<br\s*/?>', part, re.IGNORECASE):
                # Line break
                paragraph.add_run().add_break()
            else:
                # Regular text
                paragraph.add_run(part)


# =========================================================
# Convenience Functions
# =========================================================

def export_content_to_docx(content: str, title: str = "Document") -> bytes:
    """
    Quick function to export content to DOCX.

    Args:
        content: Markdown-like content
        title: Document title

    Returns:
        DOCX as bytes
    """
    exporter = DocumentExporter()
    return exporter.export_to_docx(content, title)


def export_test_plan_to_docx(test_plan: Any) -> bytes:
    """
    Quick function to export test plan to DOCX.

    Args:
        test_plan: TestPlan object

    Returns:
        DOCX as bytes
    """
    exporter = DocumentExporter()
    return exporter.export_test_plan_to_docx(test_plan)


# =========================================================
# Main (for testing)
# =========================================================

if __name__ == "__main__":
    # Test the exporter
    test_content = """
# Test Document

This is a **test** document with some content.

## Section 1

Some paragraph text here.

| Column 1 | Column 2 | Column 3 |
|----------|----------|----------|
| Value 1  | Value 2  | Value 3  |
| Value 4  | Value 5  | Value 6  |

## Section 2

More content with **bold** text.
"""

    exporter = DocumentExporter()
    docx_bytes = exporter.export_to_docx(test_content, "Test Export")

    # Save to file for testing
    with open("test_export.docx", "wb") as f:
        f.write(docx_bytes)

    print("Test document exported to test_export.docx")
